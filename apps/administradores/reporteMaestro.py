from django.http import HttpResponse
# Manejo Word
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches
#Modelos 
from apps.reportes.models import Numeral, Citas, ReporteEnviado, Modelo1, Modelo2, Modelo3, Modelo4, Modelo5, Modelo6, Modelo7, Modelo8, Modelo9, Modelo10
from apps.reportes.models import Modelo11, Modelo12, Modelo13, Modelo14, Modelo15, Modelo16, Periodo
from apps.registration.models import User
from apps.biblioteca.models import Biblioteca
from docx.shared import RGBColor

def tablaNombramiento(document):
        
        #Variables contadoras de cada Categoria 
        investigadorPosdoctoral = User.objects.filter(categoria = 'Investigador posdoctoral' ).count()
        catedraConacyt = User.objects.filter(categoria='Cátedra CONACyT').count()
        investigadorAsociadoC = User.objects.filter(categoria= 'Investigador Asociado C').count()
        investigadorTitularA = User.objects.filter(categoria= 'Investigador Titular A').count()
        investigadorTitularB = User.objects.filter(categoria = 'Investigador Titular B').count()
        investigadorTitularC = User.objects.filter(categoria = 'Investigador Titular C').count()
        investigadorTitularD = User.objects.filter(categoria = 'Investigador Titular D').count()
        TotalNombramientos = investigadorPosdoctoral + catedraConacyt + investigadorAsociadoC + investigadorTitularA + investigadorTitularB + investigadorTitularC + investigadorTitularD

        table = document.add_table(rows=1, cols=2, style='Medium Shading 1 Accent 1')
            
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nombramiento'
        hdr_cells[1].text = 'Número'
   
        row_cells = table.add_row().cells
        row_cells[0].text = "Postdocs"
        row_cells[1].text = str(investigadorPosdoctoral)
            
        row_cells = table.add_row().cells
        row_cells[0].text = "Cátedra CONACyT"
        row_cells[1].text = str(catedraConacyt)

        row_cells = table.add_row().cells
        row_cells[0].text = "Asociado C"
        row_cells[1].text = str(investigadorAsociadoC)

        row_cells = table.add_row().cells
        row_cells[0].text = "Investigador Titular A"
        row_cells[1].text = str(investigadorTitularA)

        row_cells = table.add_row().cells
        row_cells[0].text = "Investigador Titular B"
        row_cells[1].text = str(investigadorTitularB)

        row_cells = table.add_row().cells
        row_cells[0].text = "Investigador Titular C"
        row_cells[1].text = str(investigadorTitularC)

        row_cells = table.add_row().cells
        row_cells[0].text = "Investigador Titular D"
        row_cells[1].text = str(investigadorTitularD)

        row_cells = table.add_row().cells
        row_cells[0].text = "Total"
        row_cells[1].text = str(TotalNombramientos)
        document.add_paragraph()

def tablaNivelSNI(document):
    #Variables contadoras del Nivel SNI:
    sniCandidato = User.objects.filter(nivelSni='Candidato').count()
    sniNivel1 = User.objects.filter(nivelSni='Nivel 1').count()
    sniNivel2 = User.objects.filter(nivelSni='Nivel 2').count()
    sniNivel3 = User.objects.filter(nivelSni='Nivel 3').count()
    sniEmerito = User.objects.filter(nivelSni='Emérito').count()
    TotalSNI = sniCandidato + sniNivel1  + sniNivel2 + sniNivel3 + sniEmerito

    table = document.add_table(rows=1, cols=2, style='Medium Shading 1 Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER #ALINEACION CENTRAL
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Distinción SNI'
    hdr_cells[1].text = 'Número'

    row_cells = table.add_row().cells
    row_cells[0].text = "Candidato"
    row_cells[1].text = str(sniCandidato)
            
    row_cells = table.add_row().cells
    row_cells[0].text = "Nivel 1"
    row_cells[1].text = str(sniNivel1)

    row_cells = table.add_row().cells
    row_cells[0].text = "Nivel 2"
    row_cells[1].text = str(sniNivel2)

    row_cells = table.add_row().cells
    row_cells[0].text = "Nivel 3"
    row_cells[1].text = str(sniNivel3)

    row_cells = table.add_row().cells
    row_cells[0].text = "Emérito"
    row_cells[1].text = str(sniEmerito)

    row_cells = table.add_row().cells
    row_cells[0].text = "Total"
    row_cells[1].text = str(TotalSNI)

    document.add_paragraph()

def tablaInvestigadores(document,usuario):
    table = document.add_table(rows=1, cols=5, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Apellido'
    hdr_cells[1].text = 'Nombre'
    hdr_cells[2].text = 'Categoria'
    hdr_cells[3].text = 'SNI'
    hdr_cells[4].text = 'Reporte enviado'
            
    for investigador in usuario:
        try:
            reporte = ReporteEnviado.objects.get(usuario_id = investigador.id, periodo = periodo)
        except:
            reporte=""

        row_cells = table.add_row().cells
        row_cells[0].text = investigador.apellido
        row_cells[1].text = investigador.nombre
        row_cells[2].text = investigador.categoria
        row_cells[3].text = investigador.nivelSni
        if reporte:
            row_cells[4].text = 'ok'
        else:
            row_cells[4].text = ''

    document.add_page_break()

def tablaResumenNumerales(document,yearPeriodo):
    paragraph = document.add_paragraph()
    paragraph.add_run("I.Resumen: Investigación").bold = True
    table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
            
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells
    row_cells[0].text = "1"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en el primer cuartil"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,1)

    row_cells = table.add_row().cells
    row_cells[0].text = "2"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en segundo o tercer cuartil."
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,2)

    row_cells = table.add_row().cells
    row_cells[0].text = "3"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en cuarto cuartil"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,3)

    row_cells = table.add_row().cells
    row_cells[0].text = "4"
    row_cells[1].text = "Artículos científicos arbitrados en revistas del Índice CONACYT"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,4)

    row_cells = table.add_row().cells
    row_cells[0].text = "5"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas emergentes"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,5)

    row_cells = table.add_row().cells
    row_cells[0].text = "6"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas no indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo,6)

    row_cells = table.add_row().cells
    row_cells[0].text = "7"
    row_cells[1].text = "Artículos aceptados con arbitraje internacional en revistas periódicas indizadas"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,7)

    row_cells = table.add_row().cells
    row_cells[0].text = "8"
    row_cells[1].text = "Artículos aceptados con arbitraje en revistas periódicas no indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo,8)

    row_cells = table.add_row().cells
    row_cells[0].text = "9"
    row_cells[1].text = "Artículos enviados con arbitraje internacional en revistas periódicas indizadas"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,9)

    row_cells = table.add_row().cells
    row_cells[0].text = "10"
    row_cells[1].text = "Artículos enviados con arbitraje en revistas periódicas no indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo,10)

    row_cells = table.add_row().cells
    row_cells[0].text = "11"
    row_cells[1].text = "Artículos científicos arbitrados en extenso en memorias de congresos internacionales"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,11)

    row_cells = table.add_row().cells
    row_cells[0].text = "12"
    row_cells[1].text = "Artículos científicos arbitrados en extenso en memorias de congresos nacionales"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,12)

    row_cells = table.add_row().cells
    row_cells[0].text = "13"
    row_cells[1].text = "Artículos científicos no arbitrados en extenso en memorias de congresos internacionales"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,13)

    row_cells = table.add_row().cells
    row_cells[0].text = "14"
    row_cells[1].text = "Artículos científicos no arbitrados en extenso en memorias de congresos nacionales"
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,14)

    row_cells = table.add_row().cells
    row_cells[0].text = "14a"
    row_cells[1].text = "Reportes científicos no arbitrados en publicaciones periódicas."
    row_cells[2].text = Biblioteca.objects.entradasContador(yearPeriodo,15)

    row_cells = table.add_row().cells
    row_cells[0].text = "15"
    row_cells[1].text = "Autor o coautor de libros (no memorias de congreso)"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo,16)

    row_cells = table.add_row().cells
    row_cells[0].text = "16"
    row_cells[1].text = "Autor de capítulo de libro (no del mismo libro y no memoria de congreso)"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo,17)

    row_cells = table.add_row().cells
    row_cells[0].text = "17"
    row_cells[1].text = "Edición de libros / memorias"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo,18)

    row_cells = table.add_row().cells 
    row_cells[0].text = "18"
    row_cells[1].text = "Proyectos CONACyT"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo,19)

    row_cells = table.add_row().cells
    row_cells[0].text = "19"
    row_cells[1].text = "Proyectos institucionales"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo,20)

    row_cells = table.add_row().cells
    row_cells[0].text = "20"
    row_cells[1].text = "Proyectos externos"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo,21)

    row_cells = table.add_row().cells
    row_cells[0].text = "21"
    row_cells[1].text = "Proyectos interinstitucionales"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo,22)

    row_cells = table.add_row().cells
    row_cells[0].text = "22"
    row_cells[1].text = "Proyectos comercializados"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo,23)

    row_cells = table.add_row().cells
    row_cells[0].text = "23"
    row_cells[1].text = "Participación en el comité científico de conferencias internacionales (Scientific Organizing Committee; Steering Committee; similares)"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo,19)

    row_cells = table.add_row().cells
    row_cells[0].text = "24"
    row_cells[1].text = "Conferencias científicas internacionales."
    row_cells[2].text = Modelo3.objects.entradasContador(yearPeriodo,25)

    row_cells = table.add_row().cells
    row_cells[0].text = "25"
    row_cells[1].text = "Conferencias científicas nacionales"
    row_cells[2].text = Modelo3.objects.entradasContador(yearPeriodo,26)

    row_cells = table.add_row().cells
    row_cells[0].text = "26"
    row_cells[1].text = "Pláticas invitadas en conferencias internacionales"
    row_cells[2].text = Modelo3.objects.entradasContador(yearPeriodo,27)

    row_cells = table.add_row().cells
    row_cells[0].text = "27"
    row_cells[1].text = "Pláticas invitadas en conferencias nacionales"
    row_cells[2].text = Modelo3.objects.entradasContador(yearPeriodo,28)

    row_cells = table.add_row().cells
    row_cells[0].text = "28"
    row_cells[1].text = "Resúmenes en congreso internacionales"
    row_cells[2].text = Modelo3.objects.entradasContador(yearPeriodo,29)

    row_cells = table.add_row().cells
    row_cells[0].text = "29"
    row_cells[1].text = "Resúmenes en congreso nacionales"
    row_cells[2].text = Modelo3.objects.entradasContador(yearPeriodo,30)
            
    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.add_run("II.	Resumen: Formación de recursos humanos").bold = True
    table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells
    row_cells[0].text = "31"
    row_cells[1].text = "Alumnos graduados de doctorado en tiempos PNPC"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo,32)
                        
    row_cells = table.add_row().cells
    row_cells[0].text = "32"
    row_cells[1].text = "Alumnos graduados de doctorado fuera de tiempo PNPC"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo,33)
            
    row_cells = table.add_row().cells
    row_cells[0].text = "33"
    row_cells[1].text = "Alumnos graduados de maestría en tiempos PNPC"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo,34)
                        
    row_cells = table.add_row().cells
    row_cells[0].text = "34"
    row_cells[1].text = "Alumnos graduados de maestría fuera de tiempo PNPC"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo,35)
    document.add_paragraph()

    paragraph = document.add_paragraph() 
    paragraph.add_run("III.	Resumen: Desarrollo tecnológico e innovación").bold = True
    table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells # Modelo 7  40-41-42-43-44
    row_cells[0].text = "40"
    row_cells[1].text = "Derechos de autor y aseguramiento de propiedad intelectual"
    row_cells[2].text = Modelo7.objects.entradasContador(yearPeriodo,41)

    row_cells = table.add_row().cells
    row_cells[0].text = "41"
    row_cells[1].text = "Patentes solicitadas"
    row_cells[2].text = Modelo7.objects.entradasContador(yearPeriodo,42)

    row_cells = table.add_row().cells
    row_cells[0].text = "42"
    row_cells[1].text = "Patentes en proceso de evaluación que ya aprobaron el examen de forma (IMPI)"
    row_cells[2].text = Modelo7.objects.entradasContador(yearPeriodo,43)

    row_cells = table.add_row().cells
    row_cells[0].text = "44"
    row_cells[1].text = "Patentes licenciadas"
    row_cells[2].text = Modelo7.objects.entradasContador(yearPeriodo,45)

    row_cells = table.add_row().cells
    row_cells[0].text = "45"
    row_cells[1].text = "Dirección de proyectos de investigación tecnológica"
    row_cells[2].text = Modelo8.objects.entradasContador(yearPeriodo,46)

    row_cells = table.add_row().cells
    row_cells[0].text = "46"
    row_cells[1].text = "Reportes técnicos registrados"
    row_cells[2].text = Modelo9.objects.entradasContador(yearPeriodo,47)

    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.add_run("IV. Resumen: Apoyo Institucional").bold = True
    table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells
    row_cells[0].text = "48"
    row_cells[1].text = "Artículos de divulgación científica en medios masivos"
    row_cells[2].text = Modelo15.objects.entradasContador(yearPeriodo,49)

    row_cells = table.add_row().cells
    row_cells[0].text = "49"
    row_cells[1].text = "Conferencias de divulgación en eventos masivos"
    row_cells[2].text = Modelo10.objects.entradasContador(yearPeriodo,50)

    row_cells = table.add_row().cells
    row_cells[0].text = "50"
    row_cells[1].text = "Conferencias de difusión o promoción externas"
    row_cells[2].text = Modelo10.objects.entradasContador(yearPeriodo,53)
                        
    row_cells = table.add_row().cells
    row_cells[0].text = "51"
    row_cells[1].text = "Conferencias de difusión o promoción internas"
    row_cells[2].text = Modelo10.objects.entradasContador(yearPeriodo,54)
                        
    row_cells = table.add_row().cells
    row_cells[0].text = "52"
    row_cells[1].text = "Organización de eventos académicos vinculados al quehacer institucional"
    row_cells[2].text = Modelo10.objects.entradasContador(yearPeriodo,55) 

def nombreCorto(cadenaAutores,nombreC):
    nombre = nombreC.split(",")
    autores = cadenaAutores.replace(nombre[0],'<strong>{nombre}</strong>')
    return str(autores.format(nombre=nombre[0]))

def nombreFile(nombreFile):
    nombre = nombreFile.replace("anexos/","")
    return nombre

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def Reporte(request,periodo_id):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    #Objetos
    periodo = Periodo.objects.get(id=periodo_id) #Para obtener valores del 2019
    yearPeriodo = periodo.fecha_inicio.year
    fecha_inicioPeriodo = periodo.fecha_inicio
    biblioteca = Biblioteca.objects.filter(fecha_ano=yearPeriodo)
    usuario = User.objects.filter(is_staff = 0) #exclude(email = 'astrofi@inaoep.mx')
    modelo1 = Modelo1.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo2 = Modelo2.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo3 = Modelo3.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo4 = Modelo4.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo5 = Modelo5.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo6 = Modelo6.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo7 = Modelo7.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo8 = Modelo8.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo9 = Modelo9.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo10 = Modelo10.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo11 = Modelo11.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo12 = Modelo12.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo13 = Modelo13.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo14 = Modelo14.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo15 = Modelo15.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    modelo16 = Modelo16.objects.filter(periodo__fecha_inicio__year = yearPeriodo)
    citas = Citas.objects.all()
    numeral = Numeral.objects.all()
    #Fin objetos

    if fecha_inicioPeriodo.month == 1:
        MesInicio = "ENERO"
        MesFin = "JUNIO"
    elif fecha_inicioPeriodo.month == 6:
        MesInicio = "JUNIO"
        MesFin = "DICIEMBRE"

    #Titulos del documento
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    p = paragraph.add_run()
    p.bold = True
    p.text = "INAOE"
    p.font.size = Pt(13) #
            
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph.add_run()
    p.bold = True
    p.text = "Coordinación de Astrofísica"
    p.font.size = Pt(13) #

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.underline = True
    p = paragraph.add_run()
    p.bold = True
    p.text = "Reporte de Productividad Unificado"
    p.font.size = Pt(13) #
                    
    paragraph = document.add_paragraph()
    p = paragraph.add_run()
    p.text = "PERIODO DEL REPORTE: "+MesInicio + "-"+MesFin+" DE "+ str(yearPeriodo)
    p.bold = True
    p.underline = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #Fin titulos documento

    document.add_paragraph()
    paragraph = document.add_paragraph()
    p = paragraph.add_run()
    p.bold = True
    p.text = "Estadísticas de reportes: Coordinación de Astrofísica"
            
    paragraph = document.add_paragraph("Número total de investigadores en la plataforma de reportes: ")
    p = paragraph.add_run()
    p.bold = True
    p.text = str(usuario.count())

    #Tabla 1 de nombramiento 
    tablaNombramiento(document)
            
    #Tabla 2 de Distincion SNI
    tablaNivelSNI(document)

    #Tabla 3 Datos reportes
    tablaInvestigadores(document,usuario)

    #Tabla 4 resumen Numerales
    tablaResumenNumerales(document,yearPeriodo)

    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.add_run("I. INVESTIGACIÓN CIENTÍFICA").bold = True
    for n in numeral:
        
        if n.orden == 31:
            paragraph = document.add_paragraph()
            paragraph.add_run("II. FORMACIÓN DE RECURSOS HUMANOS").bold = True

        if n.orden == 40:
            paragraph = document.add_paragraph()
            paragraph.add_run("III. DESARROLLO TECNOLÓGICO E INNOVACIÓN").bold = True

        if n.orden == 47:
            paragraph = document.add_paragraph()
            paragraph.add_run("IV. APOYO INSTITUCIONAL").bold = True

        if n.orden == 61:
            paragraph = document.add_paragraph()
            paragraph.add_run("V. INFORMACIÓN ADICIONAL").bold = True
        
        document.add_paragraph(n.nombre) #Nombre - Numerales
        if n.orden == 30:
            #Tabla Distincion SNI
            table = document.add_table(rows=1, cols=4, style='Medium Shading 1 Accent 1')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER #ALINEACION CENTRAL
            table.allow_autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Investigador'
            hdr_cells[1].text = 'Citas totales'
            hdr_cells[2].text = 'Citas en el periodo'
            hdr_cells[3].text = 'Índice-h'
            for item in citas: 
                for inv in usuario:
                    if n.id == item.numeral_id and inv.id == item.usuario_id:   
                        row_cells = table.add_row().cells
                        row_cells[0].text = inv.nombreCorto
                        row_cells[1].text = str(item.citas)
                        row_cells[2].text = str(item.citas_en_periodo)
                        row_cells[3].text = str(item.indiceH)
        
        for inv in usuario:
            for item in biblioteca:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    paragraph.add_run(item.autores)
                    paragraph.add_run(",")
                    
                    paragraph.add_run(item.titulo).font.italic = True
                    paragraph.add_run(",")
                    
                    paragraph.add_run(item.revista_publicacion)
                    paragraph.add_run(" ")
                    
                    paragraph.add_run(item.paginas)
                    paragraph.add_run(" ")
                    
                    paragraph.add_run(item.volumen)
                    paragraph.add_run(" ")
                    
                    paragraph.add_run(item.fecha)
                    
                    p = document.add_paragraph() 
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Inches(0.25)

                    if item.doi:
                        p.add_run("DOI: ")
                        add_hyperlink(p,item.doi,"https://ui.adsabs.harvard.edu/link_gateway/"+item.bibcode+"/doi:"+item.doi)
                        p.add_run(" \r")
                    
                    p.add_run("Bibcode: "+item.bibcode)
                    p.add_run(" \r")
                    
                    if item.estudiantes_en_articulo:
                        p.add_run("Estudiante(s): ")
                        p.add_run(item.estudiantes_en_articulo).font.color.rgb = RGBColor(255,0,0)
                        p.add_run(" \r")
                    
                    p.add_run("URL: ")
                    add_hyperlink(p,item.url,item.url)
                    p.add_run(" \r")

                    if item.anexos:
                        p.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo1:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    paragraph.add_run(item.autores)
                    paragraph.add_run(" ")

                    paragraph.add_run(item.titulo).font.italic = True
                    paragraph.add_run(",")

                    if item.revista_publicacion:
                        paragraph.add_run(item.revista_publicacion)
                        paragraph.add_run(" ")

                    paragraph.add_run(item.fecha)
                        
                    p = document.add_paragraph() 
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                    p.paragraph_format.left_indent = Inches(0.25)

                    if item.doi:
                        p.add_run("DOI: ")
                        p.add_run(item.doi).font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                        p.add_run(" \r")
                    
                    if item.estudiantes_en_articulo:
                        p.add_run("Estudiante(s): ")
                        p.add_run(item.estudiantes_en_articulo).font.color.rgb = RGBColor(255,0,0)
                        p.add_run(" \r")
                    
                    if item.url:
                        p.add_run("URL: ")
                        add_hyperlink(p,item.url,item.url)
                        p.add_run(" \r")
                    
                    if item.anexos:
                        p.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo2:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre_del_proyecto:
                        paragraph.add_run("Nombre del proyecto: " + item.nombre_del_proyecto + "\n")

                    if item.participantes:
                        paragraph.add_run("Participantes: " + item.participantes + "\n")

                    if item.rol:
                        paragraph.add_run("Rol: " + item.rol + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripcion: " + item.descripcion + "\n")

                    if item.estudiantes:
                        paragraph.add_run("Estudiantes: ")
                        paragraph.add_run(item.estudiantes + "\n").font.color.rgb = RGBColor(255,0,0)

                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo3:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.titulo_de_la_presentacion:
                        paragraph.add_run("Título de la presentación: " + item.titulo_de_la_presentacion + "\n")

                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")

                    if item.nombre_de_conferencia:
                        paragraph.add_run("Nombre de la conferencia: " + item.nombre_de_conferencia + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.tipo:
                        paragraph.add_run("Tipo de presentación: " + item.tipo + "\n")

                    if item.estudiantes:
                        paragraph.add_run("Estudiantes: ")
                        paragraph.add_run(item.estudiantes + "\n").font.color.rgb = RGBColor(255,0,0)

                    if item.doi:
                        paragraph.add_run("DOI/ISBN: " + item.doi + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo4:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre_completo:
                        paragraph.add_run("Nombre completo: " + item.nombre_completo + "\n")

                    if item.titulo_de_tesis:
                        paragraph.add_run("Título de tesis: " + item.titulo_de_tesis + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo5:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre_del_curso:
                        paragraph.add_run("Nombre del curso: " + item.nombre_del_curso + "\n")

                    if item.periodo_numeral:
                        paragraph.add_run("Periodo: " + item.periodo_numeral + "\n")

                    if item.notas:
                        paragraph.add_run("Notas: " + item.notas + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo6:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre:
                        paragraph.add_run("Nombre: " + item.nombre + "\n")

                    if item.titulo_de_tesis:
                        paragraph.add_run("Título de tesis: " + item.titulo_de_tesis + "\n")

                    if item.grado:
                        paragraph.add_run("grado: " + item.grado + "\n")

                    if item.institucion:
                        paragraph.add_run("Institución: " + item.institucion + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.notas:
                        paragraph.add_run("Notas: " + item.notas + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo7:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo8:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre:
                        paragraph.add_run("Nombre: " + item.nombre + "\n")

                    if item.participantes:
                        paragraph.add_run("Participantes: " + item.participantes + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")

                    if item.financiamiento:
                        paragraph.add_run("Financiamiento: " + item.financiamiento + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo9:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.titulo:
                        paragraph.add_run("Titulo: " + item.titulo + "\n")
                    
                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")
                    
                    if item.numero_reportes:
                        paragraph.add_run("No. Reporte/ID: " + item.numero_reportes + "\n")
                    
                    if item.revista_publicacion:
                        paragraph.add_run("Revista o publicación: " + item.revista_publicacion + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.doi:
                        paragraph.add_run("DOI/ISBN: " + item.doi + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                    
            for item in modelo10:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    
                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.periodo_numeral:
                        paragraph.add_run("Periodo: " + item.periodo_numeral)
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url + "\n")
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                   
            for item in modelo11:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombre_del_estudiante:
                        paragraph.add_run("Nombre del estudiante: " + item.nombre_del_estudiante + "\n")
       
                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo12:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.laboratorio_taller:
                        paragraph.add_run("Laboratorio o taller:" + item.laboratorio_taller + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                   
            for item in modelo13:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.agencias_financieras:
                        paragraph.add_run("Agencia(s) financiadora(s):" + item.agencias_financieras + "\n")
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                   
            for item in modelo14:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.telescopio_instrumento_infra:
                        paragraph.add_run("Telescopio, instrumento, infraestructura: " + item.telescopio_instrumento_infra + "\n")
                    
                    if item.conferencia_proyecto:
                        paragraph.add_run("Conferencia: " + item.conferencia_proyecto + "\n")
                    
                    if item.rol:
                        paragraph.add_run("Rol: " + item.rol + "\n")
                    
                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url + "\n")
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo15:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion)
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo16:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre_del_estudiante:
                        paragraph.add_run("Nombre estudiante: " + item.nombre_del_estudiante + "\n")
                    
                    if item.coordinacion:
                        paragraph.add_run("Coordinacion: " + item.coordinacion + "\n")

                    if item.grado:
                        paragraph.add_run("grado: " + item.grado + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                   
    
            

        """
        if n.id == 1:#Header del documento
            pass
        """
            
    return document