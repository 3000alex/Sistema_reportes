from django.http import HttpResponse
# Manejo Word
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches
#Modelos 
from investigadores.models import Numeral, Citas, ReporteEnviado, Modelo1, Modelo2, Modelo3, Modelo4, Modelo5, Modelo6, Modelo7, Modelo8, Modelo9, Modelo10
from investigadores.models import Modelo11, Modelo12, Modelo13, Modelo14, Modelo15, Modelo16, Periodo
from registration.models import User
from biblioteca.models import Biblioteca
from docx.shared import RGBColor

def nombreCorto(cadenaAutores,nombreC):
    nombre = nombreC.split(",")
    autores = cadenaAutores.replace(nombre[0],'<strong>{nombre}</strong>')
    return str(autores.format(nombre=nombre[0]))

def nombreFile(nombreFile):
    nombre = nombreFile.replace("anexos/","")
    return nombre

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def Reporte(request,periodo_id):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    #Objetos
    #periodo = Periodo.objects.last()
    periodo = Periodo.objects.get(id=periodo_id) #Para obtener valores del 2019
    yearPeriodo = periodo.fechaInicio.year
    MesFin = ""
    MesInicio = ""

    if periodo.fechaInicio.month == 1:
        mesFin = "JUNIO"
    elif periodo.fechaInicio.month == 7:
        mesFin = "DICIEMBRE"

    biblioteca = Biblioteca.objects.filter(fecha_ano=yearPeriodo)

    usuario = User.objects.filter(is_staff = 0) #exclude(email = 'astrofi@inaoep.mx')
    modelo1 = Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo2 = Modelo2.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo3 = Modelo3.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo4 = Modelo4.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo5 = Modelo5.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo6 = Modelo6.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo7 = Modelo7.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo8 = Modelo8.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo9 = Modelo9.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo10 = Modelo10.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo11 = Modelo11.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo12 = Modelo12.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo13 = Modelo13.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo14 = Modelo14.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo15 = Modelo15.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    modelo16 = Modelo16.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    citas = Citas.objects.filter(periodo__fechaInicio__year = yearPeriodo)
    numeral = Numeral.objects.all()
 
    #Variables contadoras de cada Categoria 
    investigadorPosdoctoral = User.objects.filter(categoria = 'Investigador posdoctoral' ).count()
    catedraConacyt = User.objects.filter(categoria='Cátedra CONACyT').count()
    investigadorAsociadoC = User.objects.filter(categoria= 'Investigador Asociado C').count()
    investigadorTitularA = User.objects.filter(categoria= 'Investigador Titular A').count()
    investigadorTitularB = User.objects.filter(categoria = 'Investigador Titular B').count()
    investigadorTitularC = User.objects.filter(categoria = 'Investigador Titular C').count()
    investigadorTitularD = User.objects.filter(categoria = 'Investigador Titular D').count()
    TotalNombramientos = investigadorPosdoctoral + catedraConacyt + investigadorAsociadoC + investigadorTitularA + investigadorTitularB + investigadorTitularC + investigadorTitularD
    #Variables contadoras del Nivel SNI:
    sniCandidato = User.objects.filter(nivelSni='Candidato').count()
    sniNivel1 = User.objects.filter(nivelSni='Nivel 1').count()
    sniNivel2 = User.objects.filter(nivelSni='Nivel 2').count()
    sniNivel3 = User.objects.filter(nivelSni='Nivel 3').count()
    sniEmerito = User.objects.filter(nivelSni='Emérito').count()
    TotalSNI = sniCandidato + sniNivel1  + sniNivel2 + sniNivel3 + sniEmerito
    
    fechaInicioPeriodo = periodo.fechaInicio
    yearPeriodo = str(fechaInicioPeriodo.year)
    #Fin objetos

    if fechaInicioPeriodo.month == 1:
        MesInicio = "ENERO"
        MesFin = "JUNIO"
    elif fechaInicioPeriodo.month == 6:
        MesInicio = "JUNIO"
        MesFin = "DICIEMBRE"

    for n in numeral:
        if n.id == 1:#Header del documento
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            p = paragraph.add_run()
            p.bold = True
            p.text = "INAOE"
            p.font.size = Pt(13) #
            

            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = paragraph.add_run()
            p.bold = True
            p.text = "Coordinación de Astrofísica"
            p.font.size = Pt(13) #

            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.underline = True
            p = paragraph.add_run()
            p.bold = True
            p.text = "Reporte de Productividad Unificado"
            p.font.size = Pt(13) #
            
            
            paragraph = document.add_paragraph()
            
            p = paragraph.add_run()
            p.text = "PERIODO DEL REPORTE: "+MesInicio + "-"+MesFin+" DE "+yearPeriodo
            p.bold = True
            p.underline = True
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            document.add_paragraph()
            paragraph = document.add_paragraph()
            p = paragraph.add_run()
            p.bold = True
            p.text = "Estadísticas de Investigadores: Coordinación de Astrofísica"
            
            paragraph = document.add_paragraph("Número total de investigadores en la plataforma de reportes: ")
            p = paragraph.add_run()
            p.bold = True
            p.text = str(User.objects.filter(is_staff = 0).count())

            #Tabla de nombramiento 
#grid 4 enfasis 1
            table = document.add_table(rows=1, cols=2, style='Medium Shading 1 Accent 1')
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nombramiento'
            hdr_cells[1].text = 'Número'

            
            row_cells = table.add_row().cells
            row_cells[0].text = "Postdocs"
            row_cells[1].text = str(investigadorPosdoctoral)
            
            row_cells = table.add_row().cells
            row_cells[0].text = "Cátedra CONACyT"
            row_cells[1].text = str(catedraConacyt)

            row_cells = table.add_row().cells
            row_cells[0].text = "Asociado C"
            row_cells[1].text = str(investigadorAsociadoC)

            row_cells = table.add_row().cells
            row_cells[0].text = "Investigador Titular A"
            row_cells[1].text = str(investigadorTitularA)

            row_cells = table.add_row().cells
            row_cells[0].text = "Investigador Titular B"
            row_cells[1].text = str(investigadorTitularB)

            row_cells = table.add_row().cells
            row_cells[0].text = "Investigador Titular C"
            row_cells[1].text = str(investigadorTitularC)

            row_cells = table.add_row().cells
            row_cells[0].text = "Investigador Titular D"
            row_cells[1].text = str(investigadorTitularD)

            row_cells = table.add_row().cells
            row_cells[0].text = "Total"
            row_cells[1].text = str(TotalNombramientos)
            
            document.add_paragraph()

            #Tabla Distincion SNI
            table = document.add_table(rows=1, cols=2, style='Medium Shading 1 Accent 1')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER #ALINEACION CENTRAL
            table.allow_autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Distinción SNI'
            hdr_cells[1].text = 'Número'

            row_cells = table.add_row().cells
            row_cells[0].text = "Candidato"
            row_cells[1].text = str(sniCandidato)
            
            row_cells = table.add_row().cells
            row_cells[0].text = "Nivel 1"
            row_cells[1].text = str(sniNivel1)

            row_cells = table.add_row().cells
            row_cells[0].text = "Nivel 2"
            row_cells[1].text = str(sniNivel2)

            row_cells = table.add_row().cells
            row_cells[0].text = "Nivel 3"
            row_cells[1].text = str(sniNivel3)

            row_cells = table.add_row().cells
            row_cells[0].text = "Emérito"
            row_cells[1].text = str(sniEmerito)

            row_cells = table.add_row().cells
            row_cells[0].text = "Total"
            row_cells[1].text = str(TotalSNI)

            document.add_paragraph()

            #Tabla 3 Datos Investigadores
            table = document.add_table(rows=1, cols=5, style='Medium Shading 1 Accent 1')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Apellido'
            hdr_cells[1].text = 'Nombre'
            hdr_cells[2].text = 'Categoria'
            hdr_cells[3].text = 'SNI'
            hdr_cells[4].text = 'Reporte enviado'
            
            for investigador in usuario:
                try:
                    reporte = ReporteEnviado.objects.get(usuario_id = investigador.id,periodo = periodo)
                except:
                    reporte=""

                row_cells = table.add_row().cells
                row_cells[0].text = investigador.last_name
                row_cells[1].text = investigador.first_name
                row_cells[2].text = investigador.categoria
                row_cells[3].text = investigador.nivelSni
                if reporte:
                    row_cells[4].text = 'ok'
                else:
                    row_cells[4].text = ''


            document.add_page_break()

            paragraph = document.add_paragraph()
            paragraph.add_run("I.Resumen: Investigación").bold = True
            table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Numeral'
            hdr_cells[1].text = 'Concepto'
            hdr_cells[2].text = 'Total en el periodo'

            row_cells = table.add_row().cells
            row_cells[0].text = "1"
            row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en el primer cuartil"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 1).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "2"
            row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en segundo o tercer cuartil."
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 2).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "3"
            row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en cuarto cuartil"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 3).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "4"
            row_cells[1].text = "Artículos científicos arbitrados en revistas del Índice CONACYT"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 4).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "5"
            row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas emergentes"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 5).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "6"
            row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas no indizadas"
            row_cells[2].text = str( Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 6).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "7"
            row_cells[1].text = "Artículos aceptados con arbitraje internacional en revistas periódicas indizadas"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 7).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "8"
            row_cells[1].text = "Artículos aceptados con arbitraje en revistas periódicas no indizadas"
            row_cells[2].text = str( Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 8).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "9"
            row_cells[1].text = "Artículos enviados con arbitraje internacional en revistas periódicas indizadas"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 9).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "10"
            row_cells[1].text = "Artículos enviados con arbitraje en revistas periódicas no indizadas"
            row_cells[2].text = str( Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 10).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "11"
            row_cells[1].text = "Artículos científicos arbitrados en extenso en memorias de congresos internacionales"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 11).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "12"
            row_cells[1].text = "Artículos científicos arbitrados en extenso en memorias de congresos nacionales"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 12).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "13"
            row_cells[1].text = "Artículos científicos no arbitrados en extenso en memorias de congresos internacionales"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 13).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "14"
            row_cells[1].text = "Artículos científicos no arbitrados en extenso en memorias de congresos nacionales"
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 14).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "14a"
            row_cells[1].text = "Reportes científicos no arbitrados en publicaciones periódicas."
            row_cells[2].text = str( Biblioteca.objects.filter(fecha_ano=yearPeriodo, numeral_id = 15).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "15"
            row_cells[1].text = "Autor o coautor de libros (no memorias de congreso)"
            row_cells[2].text = str( Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 16).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "16"
            row_cells[1].text = "Autor de capítulo de libro (no del mismo libro y no memoria de congreso)"
            row_cells[2].text = str( Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 17).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "17"
            row_cells[1].text = "Edición de libros / memorias"
            row_cells[2].text = str( Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 18).count() )

            row_cells = table.add_row().cells 
            row_cells[0].text = "18"
            row_cells[1].text = "Proyectos CONACyT"
            row_cells[2].text = str( Modelo2.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 19).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "19"
            row_cells[1].text = "Proyectos institucionales"
            row_cells[2].text = str( Modelo2.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 20).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "20"
            row_cells[1].text = "Proyectos externos"
            row_cells[2].text = str( Modelo2.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 21).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "21"
            row_cells[1].text = "Proyectos interinstitucionales"
            row_cells[2].text = str( Modelo2.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 22).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "22"
            row_cells[1].text = "Proyectos comercializados"
            row_cells[2].text = str( Modelo2.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 23).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "23"
            row_cells[1].text = "Participación en el comité científico de conferencias internacionales (Scientific Organizing Committee; Steering Committee; similares)"
            row_cells[2].text = str( Modelo1.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 24).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "24"
            row_cells[1].text = "Conferencias científicas internacionales."
            row_cells[2].text = str( Modelo3.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 25).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "25"
            row_cells[1].text = "Conferencias científicas nacionales"
            row_cells[2].text = str( Modelo3.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 26).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "26"
            row_cells[1].text = "Pláticas invitadas en conferencias internacionales"
            row_cells[2].text = str( Modelo3.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 27).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "27"
            row_cells[1].text = "Pláticas invitadas en conferencias nacionales"
            row_cells[2].text = str( Modelo3.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 28).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "28"
            row_cells[1].text = "Resúmenes en congreso internacionales"
            row_cells[2].text = str( Modelo3.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 29).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "29"
            row_cells[1].text = "Resúmenes en congreso nacionales"
            row_cells[2].text = str( Modelo3.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 30).count() )
            

            document.add_paragraph()

            paragraph = document.add_paragraph()
            paragraph.add_run("II.	Resumen: Formación de recursos humanos").bold = True
            table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Numeral'
            hdr_cells[1].text = 'Concepto'
            hdr_cells[2].text = 'Total en el periodo'

            row_cells = table.add_row().cells
            row_cells[0].text = "31"
            row_cells[1].text = "Alumnos graduados de doctorado en tiempos PNPC"
            row_cells[2].text = str( Modelo4.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 32).count() )
                        
            row_cells = table.add_row().cells
            row_cells[0].text = "32"
            row_cells[1].text = "Alumnos graduados de doctorado fuera de tiempo PNPC"
            row_cells[2].text = str( Modelo4.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 33).count() )
            
            row_cells = table.add_row().cells
            row_cells[0].text = "33"
            row_cells[1].text = "Alumnos graduados de maestría en tiempos PNPC"
            row_cells[2].text = str( Modelo4.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 34).count() )
                        
            row_cells = table.add_row().cells
            row_cells[0].text = "34"
            row_cells[1].text = "Alumnos graduados de maestría fuera de tiempo PNPC"
            row_cells[2].text = str( Modelo4.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 35).count() )
            document.add_paragraph()

            paragraph = document.add_paragraph() 
            paragraph.add_run("III.	Resumen: Desarrollo tecnológico e innovación").bold = True
            table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Numeral'
            hdr_cells[1].text = 'Concepto'
            hdr_cells[2].text = 'Total en el periodo'

            row_cells = table.add_row().cells # Modelo 7  40-41-42-43-44
            row_cells[0].text = "40"
            row_cells[1].text = "Derechos de autor y aseguramiento de propiedad intelectual"
            row_cells[2].text = str( Modelo7.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 41).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "41"
            row_cells[1].text = "Patentes solicitadas"
            row_cells[2].text = str( Modelo7.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 42).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "42"
            row_cells[1].text = "Patentes en proceso de evaluación que ya aprobaron el examen de forma (IMPI)"
            row_cells[2].text = str( Modelo7.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 43).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "44"
            row_cells[1].text = "Patentes licenciadas"
            row_cells[2].text = str( Modelo7.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 45).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "45"
            row_cells[1].text = "Dirección de proyectos de investigación tecnológica"
            row_cells[2].text = str( Modelo8.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 46).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "46"
            row_cells[1].text = "Reportes técnicos registrados"
            row_cells[2].text = str( Modelo9.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 47).count() )

            document.add_paragraph()

            paragraph = document.add_paragraph()
            paragraph.add_run("IV. Resumen: Apoyo Institucional").bold = True
            table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 1')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Numeral'
            hdr_cells[1].text = 'Concepto'
            hdr_cells[2].text = 'Total en el periodo'

            row_cells = table.add_row().cells
            row_cells[0].text = "48"
            row_cells[1].text = "Artículos de divulgación científica en medios masivos"
            row_cells[2].text = str( Modelo15.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 49 ).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "49"
            row_cells[1].text = "Conferencias de divulgación en eventos masivos"
            row_cells[2].text = str( Modelo10.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 50).count() )

            row_cells = table.add_row().cells
            row_cells[0].text = "50"
            row_cells[1].text = "Conferencias de difusión o promoción externas"
            row_cells[2].text = str( Modelo10.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 53).count() )
                        
            row_cells = table.add_row().cells
            row_cells[0].text = "51"
            row_cells[1].text = "Conferencias de difusión o promoción internas"
            row_cells[2].text = str( Modelo10.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 54).count() )
                        
            row_cells = table.add_row().cells
            row_cells[0].text = "52"
            row_cells[1].text = "Organización de eventos académicos vinculados al quehacer institucional"
            row_cells[2].text = str( Modelo10.objects.filter(periodo__fechaInicio__year = yearPeriodo, numeral_id = 55).count() )
            

            document.add_paragraph()

            paragraph = document.add_paragraph()
            paragraph.add_run("I. INVESTIGACIÓN CIENTÍFICA").bold = True

        if n.id == 32:
            paragraph = document.add_paragraph()
            paragraph.add_run("II. FORMACIÓN DE RECURSOS HUMANOS").bold = True

        if n.id == 41:
            paragraph = document.add_paragraph()
            paragraph.add_run("III. DESARROLLO TECNOLÓGICO E INNOVACIÓN").bold = True

        if n.id == 48:
            paragraph = document.add_paragraph()
            paragraph.add_run("IV. APOYO INSTITUCIONAL").bold = True

        if n.id == 65:
            paragraph = document.add_paragraph()
            paragraph.add_run("V. INFORMACIÓN ADICIONAL").bold = True
        
        document.add_paragraph(n.nombre) #Nombre - Numerales
        if n.id == 31:
            #Tabla Distincion SNI
            table = document.add_table(rows=1, cols=4, style='Medium Shading 1 Accent 1')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER #ALINEACION CENTRAL
            table.allow_autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Investigador'
            hdr_cells[1].text = 'Citas totales'
            hdr_cells[2].text = 'Citas en el periodo'
            hdr_cells[3].text = 'Índice-h'
            for item in citas: 
                for inv in usuario:
                    if n.id == item.numeral_id and inv.id == item.usuario_id:   
                        row_cells = table.add_row().cells
                        row_cells[0].text = inv.nombreCorto
                        row_cells[1].text = item.citas
                        row_cells[2].text = item.citasObtenidasEnPeriodo
                        row_cells[3].text = item.indiceH
    
        
        for inv in usuario: 

            for item in biblioteca:
                if n.id == item.numeral_id  and inv.id == item.user_id:
                    
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    if item.autores:
                        paragraph.add_run(item.autores)
                        paragraph.add_run(",")
                    
                    if item.titulo:
                        paragraph.add_run(item.titulo).font.italic = True
                        paragraph.add_run(",")
                    
                    if item.revistaPublicacion:
                        paragraph.add_run(item.revistaPublicacion)
                        paragraph.add_run(" ")
                    
                    if item.paginas:
                        paragraph.add_run(item.paginas)
                        paragraph.add_run(" ")
                    
                    if item.volumen:
                        paragraph.add_run(item.volumen)
                        paragraph.add_run(" ")
                    
                    if item.fecha:
                        paragraph.add_run(item.fecha)
                    
                    p = document.add_paragraph() 
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Inches(0.25)

                    if item.doi:
                        p.add_run("DOI: ")
                        add_hyperlink(p,item.doi,"https://ui.adsabs.harvard.edu/link_gateway/"+item.bibcode+"/doi:"+item.doi)
                        p.add_run(" \r")
                    
                    if item.bibcode:
                        p.add_run("Bibcode: "+item.bibcode)
                        p.add_run(" \r")
                    
                    if item.estudiantesEnArticulo:
                        p.add_run("Estudiante(s): ")
                        p.add_run(item.estudiantesEnArticulo).font.color.rgb = RGBColor(255,0,0)
                        p.add_run(" \r")
                    
                    if item.url:
                        p.add_run("URL: ")
                        add_hyperlink(p,item.url,item.url)
                        p.add_run(" \r")

                    if item.anexos:
                        p.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo1:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    if item.autores:
                        paragraph.add_run(item.autores)
                        paragraph.add_run(" ")

                    if item.titulo:
                        paragraph.add_run(item.titulo).font.italic = True
                        paragraph.add_run(",")

                    if item.revistaPublicacion:
                        paragraph.add_run(item.revistaPublicacion)
                        paragraph.add_run(" ")

                    if item.fecha:
                        paragraph.add_run(item.fecha)
                        
                    p = document.add_paragraph() 
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                    p.paragraph_format.left_indent = Inches(0.25)

                    if item.doi:
                        p.add_run("DOI: ")
                        p.add_run(item.doi).font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                        p.add_run(" \r")
                    
                    if item.estudiantesEnArticulo:
                        p.add_run("Estudiante(s): ")
                        p.add_run(item.estudiantesEnArticulo).font.color.rgb = RGBColor(255,0,0)
                        p.add_run(" \r")
                    
                    if item.url:
                        p.add_run("URL: ")
                        add_hyperlink(p,item.url,item.url)
                        p.add_run(" \r")
                    
                    if item.anexos:
                        p.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo2:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombreProyecto:
                        paragraph.add_run("Nombre del proyecto: " + item.nombreProyecto + "\n")

                    if item.participantes:
                        paragraph.add_run("Participantes: " + item.participantes + "\n")

                    if item.responsableTecParticipante:
                        paragraph.add_run("Rol: ")
                        paragraph.add_run(item.responsableTecParticipante + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripcion: " + item.descripcion + "\n")

                    if item.estudiantes:
                        paragraph.add_run("Estudiantes: ")
                        paragraph.add_run(item.estudiantes + "\n").font.color.rgb = RGBColor(255,0,0)

                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
    
            for item in modelo3:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.tituloPresentacion:
                        paragraph.add_run("Título de la presentación: " + item.tituloPresentacion + "\n")

                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")

                    if item.nombreConferencia:
                        paragraph.add_run("Nombre de la conferencia: " + item.nombreConferencia + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.presentacionPoster:
                        paragraph.add_run("Tipo de presentación: " + item.presentacionPoster + "\n")

                    if item.estudiantes:
                        paragraph.add_run("Estudiantes: ")
                        paragraph.add_run(item.estudiantes + "\n").font.color.rgb = RGBColor(255,0,0)

                    if item.doi:
                        paragraph.add_run("DOI/ISBN: " + item.doi + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                                
            for item in modelo4:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombreCompleto:
                        paragraph.add_run("Nombre completo: " + item.nombreCompleto + "\n")

                    if item.tituloTesis:
                        paragraph.add_run("Título de tesis: " + item.tituloTesis + "\n")

                    if item.conferenciaProyecto:
                        paragraph.add_run("Conferencia: " + item.conferenciaProyecto + "\n")
                    
                    if item.rol:
                        paragraph.add_run("Rol: " + item.rol + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo5:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombreCurso:
                        paragraph.add_run("Nombre del curso: " + item.nombreCurso + "\n")

                    if item.periodoNumeral:
                        paragraph.add_run("Periodo: " + item.periodoNumeral + "\n")

                    if item.notas:
                        paragraph.add_run("Notas: " + item.notas + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo6:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre:
                        paragraph.add_run("Nombre: " + item.nombre + "\n")

                    if item.tituloTesis:
                        paragraph.add_run("Título de tesis: " + item.tituloTesis + "\n")

                    if item.Grado:
                        paragraph.add_run("Grado: " + item.Grado + "\n")

                    if item.institucion:
                        paragraph.add_run("Institución: " + item.institucion + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.notas:
                        paragraph.add_run("Notas: " + item.notas + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo7:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo8:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombre:
                        paragraph.add_run("Nombre: " + item.nombre + "\n")

                    if item.participantes:
                        paragraph.add_run("Participantes: " + item.participantes + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")

                    if item.financiamiento:
                        paragraph.add_run("Financiamiento: " + item.financiamiento + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo9:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.titulo:
                        paragraph.add_run("Titulo: " + item.titulo + "\n")
                    
                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")
                    
                    if item.numeroReportes:
                        paragraph.add_run("No. Reporte/ID: " + item.numeroReportes + "\n")
                    
                    if item.revistaPublicacion:
                        paragraph.add_run("Revista o publicación: " + item.revistaPublicacion + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.doi:
                        paragraph.add_run("DOI/ISBN: " + item.doi + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url)
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                    
            for item in modelo10:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    
                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.periodoNumeral:
                        paragraph.add_run("Periodo: " + item.periodoNumeral)
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph,item.url,item.url + "\n")
                        paragraph.add_run(" \r")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                   

            for item in modelo11:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombreEstudiante:
                        paragraph.add_run("Nombre del estudiante: " + item.nombreEstudiante + "\n")
       
                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))


            for item in modelo12:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.laboratorioTaller:
                        paragraph.add_run("Laboratorio o taller:" + item.laboratorioTaller + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                   

            for item in modelo13:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.agenciasFinancieras:
                        paragraph.add_run("Agencia(s) financiadora(s):" + item.agenciasFinancieras + "\n")
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
                   

            for item in modelo14:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.TelescopioInstrumentoInfra:
                        paragraph.add_run("Telescopio, instrumento, infraestructura: " + item.TelescopioInstrumentoInfra + "\n")
                    
                    if item.url:
                        add_hyperlink(paragraph,item.url,item.url + "\n")
                        paragraph.add_run(" \r")
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

            for item in modelo15:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion)
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))
            
            for item in modelo16:
                if n.id == item.numeral_id  and inv.id == item.usuario_id:
                    document.add_paragraph(inv.nombreCorto)
                    paragraph = document.add_paragraph(style='List Bullet')
                    
                    if item.nombreEstudiante:
                        paragraph.add_run("Nombre estudiante: " + item.nombreEstudiante + "\n")
                    
                    if item.coordinacion:
                        paragraph.add_run("Coordinacion: " + item.coordinacion + "\n")

                    if item.grado:
                        paragraph.add_run("Grado: " + item.grado + "\n")
                    
                    if item.anexos:
                        paragraph.add_run("Anexo: " + nombreFile(item.anexos.name))

    return document